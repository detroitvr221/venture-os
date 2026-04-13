#!/usr/bin/env python3
"""
Word Document Generator for OpenClaw Agents

Usage by agents:
  python3 /data/.openclaw/tools/generate-docx.py \
    --title "SEO Audit Report" \
    --subtitle "www.example.com" \
    --content '{"summary":"...","sections":[{"heading":"...","body":"..."}]}' \
    --output /data/.openclaw/workspace/reports/report.docx

Or via stdin:
  echo '{"title":"...","content":{...}}' | python3 /data/.openclaw/tools/generate-docx.py --stdin

Requires: pip install python-docx
"""
import json
import sys
import argparse
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# ─── Brand Colors ──────────────────────────────────────────────────────────

SKY_BLUE = RGBColor(0x4F, 0xC3, 0xF7)
GOLD = RGBColor(0xF5, 0xC5, 0x42)
DARK_BG = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY = RGBColor(0x9C, 0xA3, 0xAF)
GREEN = RGBColor(0x22, 0xC5, 0x5E)
RED = RGBColor(0xEF, 0x44, 0x44)
YELLOW = RGBColor(0xF5, 0x9E, 0x0B)


def create_report(
    title: str,
    subtitle: str = "",
    content: dict = None,
    output_path: str = "report.docx",
    score: float = None,
    report_type: str = "report",
):
    doc = Document()

    # ── Styles ──────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)

    # ── Cover / Header ──────────────────────────────────────────────────
    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NORTHBRIDGE DIGITAL")
    run.font.size = Pt(10)
    run.font.color.rgb = SKY_BLUE
    run.font.bold = True
    run.font.letter_spacing = Pt(3)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

    # Subtitle
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(14)
        run.font.color.rgb = GRAY

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datetime.now().strftime("%B %d, %Y"))
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    # Score badge
    if score is not None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(12)
        run = p.add_run(f"Score: {score}/100")
        run.font.size = Pt(24)
        run.font.bold = True
        if score >= 76:
            run.font.color.rgb = GREEN
        elif score >= 51:
            run.font.color.rgb = YELLOW
        else:
            run.font.color.rgb = RED

    doc.add_paragraph()  # spacer
    doc.add_paragraph("─" * 60)

    # ── Content ─────────────────────────────────────────────────────────
    if content:
        # Summary
        summary = content.get("summary") or content.get("executive_summary") or ""
        if summary:
            h = doc.add_heading("Executive Summary", level=1)
            h.runs[0].font.color.rgb = SKY_BLUE
            doc.add_paragraph(summary)

        # Top priorities
        priorities = content.get("top_priorities") or content.get("critical_actions") or content.get("next_steps") or []
        if priorities:
            h = doc.add_heading("Key Priorities", level=1)
            h.runs[0].font.color.rgb = SKY_BLUE
            for i, p_text in enumerate(priorities, 1):
                p = doc.add_paragraph(f"{i}. {p_text}")
                p.paragraph_format.left_indent = Cm(1)

        # Sections / Categories
        sections = content.get("sections") or content.get("categories") or []
        for section in sections:
            name = section.get("name") or section.get("heading") or "Section"
            sec_score = section.get("score")

            heading_text = name
            if sec_score is not None:
                heading_text += f" — {sec_score}/100"

            h = doc.add_heading(heading_text, level=2)
            h.runs[0].font.color.rgb = SKY_BLUE

            body = section.get("body") or section.get("description") or ""
            if body:
                doc.add_paragraph(body)

            # Findings / Issues
            findings = section.get("findings") or section.get("issues") or []
            if findings:
                table = doc.add_table(rows=1, cols=3)
                table.style = "Light Grid Accent 1"
                hdr = table.rows[0].cells
                hdr[0].text = "Severity"
                hdr[1].text = "Finding"
                hdr[2].text = "Recommendation"

                for finding in findings:
                    row = table.add_row().cells
                    sev = finding.get("severity", "info")
                    row[0].text = sev.upper()
                    row[1].text = finding.get("title", "") + ("\n" + finding.get("description", "") if finding.get("description") else "")
                    row[2].text = finding.get("recommendation") or finding.get("fix") or finding.get("detail") or ""

        # Recommended services (for intake analysis)
        services = content.get("recommended_services") or []
        if services:
            h = doc.add_heading("Recommended Services", level=1)
            h.runs[0].font.color.rgb = GOLD
            for svc in services:
                doc.add_paragraph(f"• {svc}", style="List Bullet")

        tier = content.get("recommended_tier")
        value = content.get("estimated_monthly_value")
        if tier or value:
            h = doc.add_heading("Recommendation", level=2)
            h.runs[0].font.color.rgb = GOLD
            if tier:
                doc.add_paragraph(f"Recommended Package: {tier}")
            if value:
                doc.add_paragraph(f"Estimated Monthly Value: ${value}/mo")

        # Risk factors / Opportunities
        for key, label in [("risk_factors", "Risk Factors"), ("opportunities", "Opportunities")]:
            items = content.get(key) or []
            if items:
                h = doc.add_heading(label, level=2)
                for item in items:
                    doc.add_paragraph(f"• {item}", style="List Bullet")

        # Raw text fallback
        raw = content.get("raw_response") or ""
        if raw and not sections and not summary:
            h = doc.add_heading("Full Report", level=1)
            h.runs[0].font.color.rgb = SKY_BLUE
            doc.add_paragraph(raw)

    # ── Footer ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph("─" * 60)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generated by Northbridge Digital")
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run = p.add_run(f"\n{datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    run = p.add_run("\nwww.thenorthbridgemi.com")
    run.font.size = Pt(8)
    run.font.color.rgb = SKY_BLUE

    # ── Save ────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"DOCX saved: {output_path}")
    return output_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate branded Word documents")
    parser.add_argument("--title", default="Report")
    parser.add_argument("--subtitle", default="")
    parser.add_argument("--content", help="JSON string with report content")
    parser.add_argument("--score", type=float, default=None)
    parser.add_argument("--type", default="report")
    parser.add_argument("--output", default="/data/.openclaw/workspace/reports/report.docx")
    parser.add_argument("--stdin", action="store_true", help="Read JSON from stdin")
    args = parser.parse_args()

    if args.stdin:
        data = json.load(sys.stdin)
        create_report(
            title=data.get("title", "Report"),
            subtitle=data.get("subtitle", ""),
            content=data.get("content", {}),
            output_path=data.get("output", args.output),
            score=data.get("score"),
            report_type=data.get("type", "report"),
        )
    else:
        content = json.loads(args.content) if args.content else {}
        create_report(
            title=args.title,
            subtitle=args.subtitle,
            content=content,
            output_path=args.output,
            score=args.score,
            report_type=args.type,
        )
